"""
AI Code Guardian — Document Loader
=====================================
Reads every supported file type and returns chunked Document objects
ready for embedding and indexing.

Supported formats
-----------------
Text / markup : .md .txt .rst .yaml .yml .json .xml .csv .sql .sh
Code          : .java .py .js .ts .go .rs .kt .cpp .c .cs
Config        : Dockerfile Jenkinsfile .tf Kubernetes YAML
Rich          : .pdf (via pdfminer) .docx (via python-docx) — loaded lazily
"""
from __future__ import annotations

import hashlib
import json
import logging
import os
import re
from pathlib import Path
from typing import Iterator

from guardian.ai.config import AssistantConfig
from guardian.ai.models import Document, DocumentType

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Extension → DocumentType mapping
# ---------------------------------------------------------------------------
_EXT_TO_TYPE: dict[str, DocumentType] = {
    ".md":         DocumentType.GENERAL,
    ".txt":        DocumentType.GENERAL,
    ".rst":        DocumentType.GENERAL,
    ".json":       DocumentType.GENERAL,
    ".yaml":       DocumentType.KUBERNETES,
    ".yml":        DocumentType.KUBERNETES,
    ".xml":        DocumentType.GENERAL,
    ".csv":        DocumentType.GENERAL,
    ".sql":        DocumentType.GENERAL,
    ".sh":         DocumentType.GENERAL,
    ".tf":         DocumentType.ARCHITECTURE,
    ".java":       DocumentType.SOURCE_CODE,
    ".py":         DocumentType.SOURCE_CODE,
    ".js":         DocumentType.SOURCE_CODE,
    ".ts":         DocumentType.SOURCE_CODE,
    ".go":         DocumentType.SOURCE_CODE,
    ".rs":         DocumentType.SOURCE_CODE,
    ".kt":         DocumentType.SOURCE_CODE,
    ".cpp":        DocumentType.SOURCE_CODE,
    ".c":          DocumentType.SOURCE_CODE,
    ".cs":         DocumentType.SOURCE_CODE,
    ".php":        DocumentType.SOURCE_CODE,
    ".rb":         DocumentType.SOURCE_CODE,
    ".scala":      DocumentType.SOURCE_CODE,
    ".pdf":        DocumentType.GENERAL,
    ".docx":       DocumentType.GENERAL,
}
_FILENAME_TO_TYPE: dict[str, DocumentType] = {
    "dockerfile":  DocumentType.DOCKERFILE,
    "jenkinsfile": DocumentType.CI_CD,
    "makefile":    DocumentType.GENERAL,
    "readme":      DocumentType.README,
    "readme.md":   DocumentType.README,
}
_PATH_HINTS: dict[str, DocumentType] = {
    "owasp":    DocumentType.OWASP,
    "nist":     DocumentType.NIST,
    "cwe":      DocumentType.CWE,
    "policy":   DocumentType.SECURITY_POLICY,
    "security": DocumentType.SECURITY_POLICY,
    "arch":     DocumentType.ARCHITECTURE,
    "openapi":  DocumentType.API_SPEC,
    "swagger":  DocumentType.API_SPEC,
    "require":  DocumentType.REQUIREMENT,
    "jira":     DocumentType.REQUIREMENT,
    "brd":      DocumentType.REQUIREMENT,
    "srs":      DocumentType.REQUIREMENT,
    "report":   DocumentType.SCAN_REPORT,
    "quantum":  DocumentType.QUANTUM_REPORT,
    "risk":     DocumentType.RISK_REPORT,
    "business_intent": DocumentType.BUSINESS_INTENT,
}
_CODE_LANG_MAP: dict[str, str] = {
    ".java": "java", ".py": "python", ".js": "javascript", ".ts": "typescript",
    ".go": "go", ".rs": "rust", ".kt": "kotlin", ".cpp": "cpp",
    ".c": "c", ".cs": "csharp", ".php": "php", ".rb": "ruby", ".scala": "scala",
}


def _doc_type_for(path: Path) -> DocumentType:
    """Infer DocumentType from filename and path hints."""
    name_lower = path.name.lower()
    if name_lower in _FILENAME_TO_TYPE:
        return _FILENAME_TO_TYPE[name_lower]
    ext = path.suffix.lower()
    # Check path components for content hints
    path_str = str(path).lower()
    for hint, dtype in _PATH_HINTS.items():
        if hint in path_str:
            return dtype
    return _EXT_TO_TYPE.get(ext, DocumentType.GENERAL)


def _make_doc_id(source_path: str, chunk_index: int) -> str:
    basis = f"{source_path}::{chunk_index}"
    return hashlib.sha256(basis.encode()).hexdigest()[:16]


# ---------------------------------------------------------------------------
# Chunkers
# ---------------------------------------------------------------------------

def _chunk_text(
    text: str,
    source_path: str,
    doc_type: DocumentType,
    language: str | None,
    chunk_size: int,
    chunk_overlap: int,
) -> list[Document]:
    """Split plain text into overlapping chunks."""
    text = text.strip()
    if not text:
        return []
    chunks: list[Document] = []
    start = 0
    idx = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        # Try to break on a newline rather than mid-word
        if end < len(text):
            last_nl = chunk.rfind("\n")
            if last_nl > chunk_size // 2:
                end = start + last_nl + 1
                chunk = text[start:end]
        # Rough line estimation
        lines_before = text[:start].count("\n") + 1
        lines_after  = lines_before + chunk.count("\n")
        chunks.append(Document(
            doc_id=_make_doc_id(source_path, idx),
            content=chunk,
            source_path=source_path,
            doc_type=doc_type,
            language=language,
            start_line=lines_before,
            end_line=lines_after,
        ))
        start = end - chunk_overlap
        idx += 1
    return chunks


def _chunk_code(
    text: str,
    source_path: str,
    doc_type: DocumentType,
    language: str,
    chunk_lines: int,
    overlap_lines: int,
) -> list[Document]:
    """Split source code into line-based overlapping chunks."""
    lines = text.splitlines()
    if not lines:
        return []
    chunks: list[Document] = []
    idx = 0
    start = 0
    while start < len(lines):
        end = min(start + chunk_lines, len(lines))
        chunk_text = "\n".join(lines[start:end])
        chunks.append(Document(
            doc_id=_make_doc_id(source_path, idx),
            content=chunk_text,
            source_path=source_path,
            doc_type=doc_type,
            language=language,
            start_line=start + 1,
            end_line=end,
        ))
        start = end - overlap_lines
        if start < 0:
            start = 0
        if end == len(lines):
            break
        idx += 1
    return chunks


# ---------------------------------------------------------------------------
# Per-format readers
# ---------------------------------------------------------------------------

def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_json(path: Path) -> str:
    try:
        data = json.loads(path.read_bytes())
        return json.dumps(data, indent=2)
    except Exception:
        return _read_text(path)


def _read_pdf(path: Path) -> str:
    try:
        from pdfminer.high_level import extract_text  # type: ignore
        return extract_text(str(path))
    except ImportError:
        logger.warning("pdfminer not installed — skipping %s. pip install pdfminer.six", path)
        return ""
    except Exception as exc:
        logger.warning("PDF read failed %s: %s", path, exc)
        return ""


def _read_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        logger.warning("python-docx not installed — skipping %s. pip install python-docx", path)
        return ""
    except Exception as exc:
        logger.warning("DOCX read failed %s: %s", path, exc)
        return ""


_FORMAT_READERS = {
    ".json": _read_json,
    ".pdf":  _read_pdf,
    ".docx": _read_docx,
}


# ---------------------------------------------------------------------------
# Main loader
# ---------------------------------------------------------------------------

class DocumentLoader:
    """
    Walks a directory tree (or processes a list of uploaded file paths)
    and yields Document chunks ready for embedding.
    """

    _SKIP_DIRS = {
        ".git", ".venv", "venv", "node_modules", "__pycache__",
        "target", "build", "dist", ".idea", ".vs", ".pytest_cache",
    }

    def __init__(self, config: AssistantConfig):
        self._cfg = config

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def load_directory(self, root: Path) -> list[Document]:
        """Load all supported files under `root`."""
        docs: list[Document] = []
        for path in self._walk(root):
            try:
                docs.extend(self._load_file(path, root))
            except Exception as exc:
                logger.warning("Skipping %s: %s", path, exc)
        logger.info("Loaded %d document chunks from %s", len(docs), root)
        return docs

    def load_files(self, paths: list[Path], root: Path | None = None) -> list[Document]:
        """Load a specific list of files."""
        docs: list[Document] = []
        for path in paths:
            try:
                docs.extend(self._load_file(path, root or path.parent))
            except Exception as exc:
                logger.warning("Skipping %s: %s", path, exc)
        return docs

    def load_text(
        self,
        text: str,
        source_name: str,
        doc_type: DocumentType = DocumentType.GENERAL,
    ) -> list[Document]:
        """Load raw text directly (e.g. scan report JSON string)."""
        return _chunk_text(
            text, source_name, doc_type, None,
            self._cfg.chunk_size, self._cfg.chunk_overlap,
        )

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _walk(self, root: Path) -> Iterator[Path]:
        for dirpath, dirnames, filenames in os.walk(root):
            dirnames[:] = [d for d in dirnames if d not in self._SKIP_DIRS]
            for fn in filenames:
                fp = Path(dirpath) / fn
                if self._is_supported(fp):
                    yield fp

    def _is_supported(self, path: Path) -> bool:
        name_lower = path.name.lower()
        if name_lower in _FILENAME_TO_TYPE:
            return True
        return path.suffix.lower() in self._cfg.all_indexable_extensions

    def _load_file(self, path: Path, root: Path) -> list[Document]:
        """Read one file and return its chunks."""
        try:
            rel = str(path.relative_to(root))
        except ValueError:
            rel = str(path)

        ext       = path.suffix.lower()
        doc_type  = _doc_type_for(path)
        language  = _CODE_LANG_MAP.get(ext)
        reader    = _FORMAT_READERS.get(ext, _read_text)
        text      = reader(path)

        if not text.strip():
            return []

        if doc_type == DocumentType.SOURCE_CODE and language:
            return _chunk_code(
                text, rel, doc_type, language,
                self._cfg.code_chunk_lines,
                self._cfg.code_chunk_overlap_lines,
            )
        return _chunk_text(
            text, rel, doc_type, language,
            self._cfg.chunk_size, self._cfg.chunk_overlap,
        )
